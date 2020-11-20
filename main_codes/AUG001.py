#rule(aug001)authrgroup
from bs4 import BeautifulSoup
import re
import docx
import logging


def aug001(file_xml,file_docx):
    err_rule=[]
    error=[]

    count = 0
    lst2=[]

    try:

        with open (file_xml,'r',encoding ='utf-8') as f:
            contents = f.read()
            soup1 = BeautifulSoup(contents,"lxml")
            for i in soup1.findAll("ce:author"):
                for j in i.findAll('ce:given-name'):
                    lst2.append(j.text)
                    for k in i.findAll('ce:surname'):
                        lst2.append(k.text)
                    

                for j in soup1.findAll('sa:affiliation'):
                    lst2.append(j.text)
            #print(lst2)
            lst2_reg = ' '.join(lst2)
            reg = re.sub('[^a-zA-Z0-9]+',' ', lst2_reg)
            #print(reg)
            
            doc = docx.Document(file_docx)
            #print (len(doc.paragraphs))
            lst=[]
            for i in range(2,8):
                h = doc.paragraphs[i].text
                lst.append(h)
            lst_reg = ' '.join(lst)
            reg2 = re.sub('[^a-zA-Z0-9]+',' ', lst_reg)
            #print(reg2)
            if reg2.find(reg):
                pass
            else:
                error.append([str(contents.index(reg2)),"name and affiliation not in same order"])
            if error!=[]:
                for i in error:
                    rule_file=[]
                    rule_file.append("AUG001")
                    rule_file.append('Author group')
                    rule_file.append("The order of both authors and affiliations, as supplied in the original should never be changed")
                    rule_file.append("error")
                    rule_file.append(i[0])
                    rule_file.append(i[1])
                    err_rule.append(rule_file)


            else:
                 rule_file=[]
                 rule_file.append("AUG001")
                 rule_file.append("Author group")
                 rule_file.append('The order of both authors and affiliations, as supplied in the original should never be changed')
                 rule_file.append("no error")
                 err_rule.append(rule_file)
            #print('error')
        #      for_demo2=[]
        #      for_demo2.append("AUG001")
        #      for_demo2.append("")
        #      for_demo2.append("jid")
        #      for_demo2.append(str(contents.index('ce:author-group')))
        #      for_demo2.append("error")
        #      for_demo2.append("name affiliation order not same in docs and xml")
        #      err_rule.append(for_demo2)
    except Exception as e:
        rule_file=[]
        rule_file.append("AUG001")
        rule_file.append("Author group")
        rule_file.append('The order of both authors and affiliations, as supplied in the original should never be changed')
        rule_file.append("no error")
        err_rule.append(rule_file)
        logging.info('='*50)
        logging.exception("Got exception on main handler in AUG001 : Author group " )
        logging.shutdown()


    return err_rule
# file_jss =r'C:/Users/Digiscape/Desktop/sindhu/MNT_ELSEVIER_JOURNAL_APCATA_17035_110/APCATA_110.docx'
# file_xml =r'C:/Users/Digiscape/Desktop/sindhu/MNT_ELSEVIER_JOURNAL_APCATA_17035_110/tx1.xml'
# print(aug001(file_xml,file_jss))
